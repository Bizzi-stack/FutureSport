import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_pmc_architecture_doc():
    doc = docx.Document()
    
    # Page setup - Margins 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles & Colors
    NAVY = RGBColor(0, 38, 127)     # #00267F BFA Navy
    GOLD = RGBColor(220, 160, 0)    # BFA Gold
    DARK = RGBColor(15, 23, 42)     # Dark Slate
    GRAY = RGBColor(71, 85, 105)    # Slate Gray

    # Title Banner / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sub = p_title.add_run("TECHNICAL INTEGRATION & ARCHITECTURE SPECIFICATION\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = GOLD

    run_title = p_title.add_run("The Prime Minister's Cup — Analytical Hub & API Integration Architecture")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("Target Portal: theprimeministerscups.com  |  Engine: FutureSport / EduData Platform")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Summary
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Executive Summary")
    run_h1.font.name = "Arial"
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = NAVY

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.25
    p_exec.paragraph_format.space_after = Pt(12)
    p_exec.add_run(
        "This architectural specification details how the FutureSport Dashboard operates as the "
        "official behind-the-scenes analytical engine, match verification hub, and data pipeline "
        "for the Prime Minister's Cup (theprimeministerscups.com).\n\n"
        "By establishing a direct bi-directional synchronization layer, match data flows seamlessly "
        "from on-pitch officials (coaches, statisticians, referees) through automated commissioner verification "
        "and into the public Prime Minister's Cup tournament website in real time."
    )

    # Section 2: Match Official Data Lifecycle
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("2. Match Official Data Lifecycle & Workflows")
    run_h2.font.name = "Arial"
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = NAVY

    table_lifecycle = doc.add_table(rows=6, cols=3)
    table_lifecycle.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_lifecycle.autofit = False

    headers = ["Stage", "Responsible Role", "Action & Output"]
    for i, h_text in enumerate(headers):
        cell = table_lifecycle.rows[0].cells[i]
        set_cell_background(cell, "00267F")
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Arial"

    data_lifecycle = [
        ("1. Pre-Match Setup", "Team Coaches", "Submit starting XI, tactical formations (4-3-3, 4-4-2, 4-2-3-1), and substitutes bench."),
        ("2. Real-Time Logging", "Statistician / Live Data Entry", "Log live match clock, spatial shot maps (x,y), goals, assists, cards, and goalkeeper saves."),
        ("3. Officiating Audit", "Referee & 4th Official", "Independently log disciplinary cards and submit official post-match Referee Report."),
        ("4. Discrepancy Engine", "Match Commissioner", "Automated discrepancy engine compares Statistician vs. Referee logs. Highlights goal/card mismatches for final approval."),
        ("5. Public Broadcast", "FutureSport API Sync Core", "Verified match results push to theprimeministerscups.com via Firebase Cloud Firestore & REST endpoints.")
    ]

    for row_idx, row_data in enumerate(data_lifecycle, start=1):
        row_cells = table_lifecycle.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row_cells[col_idx]
            set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section 3: Target Architecture (theprimeministerscups.com)
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("3. Target System Architecture (theprimeministerscups.com)")
    run_h3.font.name = "Arial"
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = NAVY

    p_target = doc.add_paragraph()
    p_target.paragraph_format.line_spacing = 1.25
    p_target.paragraph_format.space_after = Pt(12)
    p_target.add_run(
        "Analysis of the live target portal (theprimeministerscups.com) confirms it relies on "
        "Firebase Cloud (Firestore & Auth) with standardized cloud collections. "
        "FutureSport's integration engine (pmcSyncEngine.js) interfaces directly with these schemas:"
    )

    table_schema = doc.add_table(rows=5, cols=3)
    table_schema.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_schema = ["Firestore Collection", "Target Data Fields", "FutureSport Source Component"]
    for i, h_text in enumerate(headers_schema):
        cell = table_schema.rows[0].cells[i]
        set_cell_background(cell, "00267F")
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Arial"

    data_schema = [
        ("matches", "id, tournamentId, stage, homeTeamId, awayTeamId, homeScore, awayScore, status, period, elapsed, kickoffTime, venue", "LiveMatch.jsx & MatchCentre.jsx"),
        ("matchEvents", "id, matchId, type (goal, yellow, red, sub, gkSave), minute, playerId, teamId, goalType, assistingPlayerId", "LiveShotModal.jsx & LiveGkSaveModal.jsx"),
        ("standings", "teamId, played, won, drawn, lost, goalsFor, goalsAgainst, points", "LeagueTable.jsx"),
        ("players", "id, teamId, name, jerseyNumber, position, documents", "SchoolPlayerRegistration.jsx")
    ]

    for row_idx, row_data in enumerate(data_schema, start=1):
        row_cells = table_schema.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            cell = row_cells[col_idx]
            set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Section 4: Implementation & Deployment Roadmap
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run("4. Implementation & Synchronization Engine Roadmap")
    run_h4.font.name = "Arial"
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = NAVY

    roadmap_points = [
        ("PMC Sync Translator (pmcSyncEngine.js): ", "Translates FutureSport internal match states into the precise JSON payload formats required by theprimeministerscups.com."),
        ("Tournament Switcher Header (App.jsx): ", "Adds a top navigation toggle ([ 🏆 Prime Minister's Cup | ⚽ National Secondary Schools League ]) with BFA Navy (#00267F) and Gold (#FFC726) branding."),
        ("Automated Commissioner Webhook (CommissionerDashboard.jsx): ", "Triggers automated data push upon official match commissioner sign-off."),
        ("Live Preview & Surge Production: ", "App published at footballplatformdev.surge.sh and version controlled at github.com/Bizzi-stack/FutureSport.")
    ]

    for title, desc in roadmap_points:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(6)
        r_t = p_bullet.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.bold = True
        r_t.font.color.rgb = NAVY
        r_d = p_bullet.add_run(desc)
        r_d.font.name = "Arial"
        r_d.font.color.rgb = DARK

    # Save document
    filename = "Prime_Ministers_Cup_Integration_Architecture.docx"
    doc.save(filename)
    print(f"Document successfully created: {filename}")

if __name__ == "__main__":
    create_pmc_architecture_doc()
