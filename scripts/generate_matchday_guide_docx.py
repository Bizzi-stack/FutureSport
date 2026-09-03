import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text_list, title="NOTE", fill_hex="F1F5F9", border_color="00267F"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=200)
    
    # Left thick border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title.upper()}\n")
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(0, 38, 127) # Navy
    
    for item in text_list:
        run_item = p.add_run(item + "\n")
        run_item.font.name = 'Calibri'
        run_item.font.size = Pt(10)
        run_item.font.color.rgb = RGBColor(51, 65, 85)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(6)

def create_document():
    doc = Document()
    
    # Page setup - 0.8 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Palette
    NAVY = RGBColor(0, 38, 127)      # #00267F
    GOLD = RGBColor(197, 145, 0)     # #C59100
    SLATE = RGBColor(30, 41, 59)     # #1E293B
    MUTED = RGBColor(100, 116, 139)  # #64748B
    EMERALD = RGBColor(5, 150, 105)  # #059669

    # Document Header Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(2)
    
    run_subtitle = title_p.add_run("PRIME MINISTER'S CUP 2026 · OFFICIAL OPERATIONAL PLAYBOOK\n")
    run_subtitle.font.name = 'Calibri'
    run_subtitle.font.size = Pt(10)
    run_subtitle.font.bold = True
    run_subtitle.font.color.rgb = GOLD
    
    run_title = title_p.add_run("Matchday Data Capture & Operations Guide\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    run_desc = title_p.add_run("Standard Operating Procedure: Squad Submission to Live Event Recording\n")
    run_desc.font.name = 'Calibri'
    run_desc.font.size = Pt(12)
    run_desc.font.color.rgb = MUTED

    # Meta Table (Author, Target, Platform, Date)
    meta_tbl = doc.add_table(rows=2, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Target Audience:", " Johnathan Cumberbatch & Field Data Capture Team"), ("System Platform:", " EduData Platform (edudata-pmcup-app.surge.sh)")],
        [("Tournament:", " Prime Minister's Cup (PMC 2026)"), ("Version & Status:", " v1.4 Pilot Production Ready")]
    ]
    for r_idx, row in enumerate(meta_data):
        for c_idx, (label, val) in enumerate(row):
            cell = meta_tbl.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl_run = p.add_run(label)
            lbl_run.bold = True
            lbl_run.font.size = Pt(9.5)
            lbl_run.font.color.rgb = SLATE
            val_run = p.add_run(val)
            val_run.font.size = Pt(9.5)
            val_run.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 1: EXECUTIVE WORKFLOW OVERVIEW
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    r = h1.add_run("1. Executive Workflow Overview")
    r.font.name = 'Arial'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p_intro = doc.add_paragraph(
        "This playbook outlines the rigorous 5-stage matchday sequence designed to eliminate data discrepancies, "
        "enforce squad roster compliance, and deliver instant live match statistics to the official Prime Minister's Cup website. "
        "Every matchday follows a synchronized hand-off across four primary stakeholders:"
    )
    p_intro.paragraph_format.space_after = Pt(8)

    # 4 Roles Summary Table
    role_tbl = doc.add_table(rows=5, cols=3)
    role_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    role_headers = ["Operational Role", "Key Responsibilities", "Primary Portal / Device"]
    for i, h_text in enumerate(role_headers):
        cell = role_tbl.cell(0, i)
        set_cell_background(cell, "00267F")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    roles_info = [
        ("⚽ Team Coach / Manager", "Selects starting XI, jersey kit numbers, and tactical formations on pitch.", "Coach Matchday Selection"),
        ("🛡️ Super-Administrator", "Validates official team sheets and audits Player IDs prior to kick-off.", "Competition Admin (Squad Hub)"),
        ("🧑‍⚖️ Match Referee", "Controls match countdown, blows kick-off whistle, verifies final scores.", "Referee Field Terminal"),
        ("📊 Data Capture Team (Johnathan)", "Records real-time possession, shots, saves, and disciplinary events.", "Statistician Dashboard")
    ]
    for r_idx, (role_name, resp, portal) in enumerate(roles_info, start=1):
        bg_color = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, val in enumerate([role_name, resp, portal]):
            cell = role_tbl.cell(r_idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if c_idx == 0:
                run.bold = True
                run.font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # -------------------------------------------------------------
    # SECTION 2: END-TO-END 5-STAGE OPERATIONAL LIFECYCLE
    # -------------------------------------------------------------
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. The 5-Stage Matchday Operational Lifecycle")
    r2.font.name = 'Arial'
    r2.font.size = Pt(15)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    stages = [
        ("STAGE 1: COACH SQUAD SELECTION & SUBMISSION", [
            "Coaches access the tactical squad selection tool on their portal.",
            "Choose team formation (e.g. 4-3-3, 4-2-3-1, 3-5-2).",
            "Assign starting XI players to tactical pitch positions with official 2D jersey kit icons.",
            "Review verified permanent Player IDs (PID-PMC-XXXXX) to ensure roster compliance.",
            "Click '🔒 Lock In & Submit Official Squad'. Roster enters status: PENDING_VALIDATION."
        ]),
        ("STAGE 2: SUPER-ADMINISTRATOR SQUAD VERIFICATION", [
            "League Super-Administrator opens the 🛡️ Squad Validation Hub in Competition Admin.",
            "Conducts side-by-side verification of Home and Away team sheets.",
            "System cross-checks against registered databases to flag duplicate entries or ineligible players.",
            "Administrator clicks '✓ Approve Squad'.",
            "Result: Squad is locked as the immutable Official Matchday Source."
        ]),
        ("STAGE 3: AUTOMATED DISPATCH & MULTI-CHANNEL ALERTS", [
            "Upon approval, the automated dispatch relay (refereeNotificationService) triggers immediately.",
            "Referee Notification: Referee receives whistle alert and confirmed lineups.",
            "Analyst Notification: Johnathan Cumberbatch receives an automated high-priority email.",
            "Email Contents: Lineup team sheets, venue, kickoff time, and a 1-Tap Direct Deep Link.",
            "Device Push: Browser chime sound and mobile lockscreen alert are delivered."
        ]),
        ("STAGE 4: ANALYST PORTAL INITIALIZATION & ROLE SCOPING", [
            "Analyst clicks direct deep link or logs into Statistician Portal (johnathan).",
            "System initializes the Tile Data Capture Control Panel.",
            "Role Scoping Matrix engages to prevent multi-user event collisions (Possession vs. Shots vs. General).",
            "Master Lead Analysts retain full access across all 13 direct event tiles with quick-switch role pills."
        ]),
        ("STAGE 5: REFEREE WHISTLE & FIRST EVENT LOGGING", [
            "Match Referee blows the whistle (▶️ Kick Off). Official Web Audio whistle synthesizes.",
            "Match clock commences live from 00:00. Match state shifts to LIVE 1H.",
            "First Event A (Possession): Logger taps the team in possession; live clock starts recording seconds.",
            "First Event B (Shot on Target): Logger taps '⚽ Shot on Target' -> selects player jersey -> taps pitch coordinates.",
            "Opposing Goalkeeper automatically credited with +1 Save in real-time.",
            "Events instantly stream to the live timeline feed and push to the public Supabase REST API."
        ])
    ]

    for st_title, st_steps in stages:
        h_st = doc.add_heading(level=2)
        r_st = h_st.add_run(st_title)
        r_st.font.name = 'Arial'
        r_st.font.size = Pt(12)
        r_st.font.bold = True
        r_st.font.color.rgb = GOLD
        h_st.paragraph_format.space_before = Pt(10)
        h_st.paragraph_format.space_after = Pt(4)

        for step in st_steps:
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.paragraph_format.space_before = Pt(1)
            p_bullet.paragraph_format.space_after = Pt(2)
            r_bullet = p_bullet.add_run(step)
            r_bullet.font.name = 'Calibri'
            r_bullet.font.size = Pt(10)
            r_bullet.font.color.rgb = SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # SECTION 3: 6-CAPTURER ROLE SCOPING SYNCHRONIZATION MATRIX
    # -------------------------------------------------------------
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. 6-Capturer Concurrent Role Scoping Matrix")
    r3.font.name = 'Arial'
    r3.font.size = Pt(15)
    r3.font.bold = True
    r3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)

    p_matrix_desc = doc.add_paragraph(
        "To guarantee zero event duplication when 6 data loggers operate concurrently during high-stakes matchdays, "
        "the platform automatically partitions logging tiles into dedicated scopes. Non-assigned tiles are greyed out "
        "with disabled click listeners:"
    )
    p_matrix_desc.paragraph_format.space_after = Pt(8)

    scope_tbl = doc.add_table(rows=5, cols=4)
    scope_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_headers = ["Assigned Scope", "Logger Count", "Interactive / Unlocked Tiles", "Disabled / Greyed Out"]
    for i, h_text in enumerate(scope_headers):
        cell = scope_tbl.cell(0, i)
        set_cell_background(cell, "00267F")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    scopes_data = [
        ("⏱️ Possession Specialist", "1 Logger", "Home / Away Ball Possession Toggles & Live Clock Switchers", "🔒 All Shot Tiles & General Event Tiles"),
        ("⚽ Shot Specialist", "2 Loggers", "Goals, Shots on Target, Shots Off Target, Blocked Shots, Penalties", "🔒 Possession Timers & General Event Tiles"),
        ("📋 General Events Specialist", "3 Loggers", "Fouls, Yellow / Red Cards, Corners, Offsides, Saves, Substitutions", "🔒 Possession Timers & Shot Tiles"),
        ("👑 Master Lead Analyst (Johnathan)", "1-2 Leads", "ALL 13 TILES FULLY UNLOCKED (Global Override Capability)", "None (Full Administrative Access)")
    ]

    for r_idx, (role_name, count, unl, dis) in enumerate(scopes_data, start=1):
        bg_color = "FFFFFF" if r_idx % 2 != 0 else "F8FAFC"
        for c_idx, val in enumerate([role_name, count, unl, dis]):
            cell = scope_tbl.cell(r_idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True
                run.font.color.rgb = NAVY
            elif c_idx == 3 and "🔒" in val:
                run.font.color.rgb = RGBColor(185, 28, 28) # Red alert color

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Callout Box for Johnathan
    add_callout_box(
        doc,
        [
            "Account: Johnathan Cumberbatch (johnathan)",
            "Email: johnathan.cumberbatch@gmail.com | Password: password",
            "Assigned Scope: Master Lead Analyst (captureRole = 'all')",
            "Johnathan can switch into any specialist scope on the fly using the top-bar pills (⏱️ Possession, ⚽ Shots, 📋 General) to inspect what junior capturers see!"
        ],
        title="Field Credentials & Master Permissions",
        fill_hex="EFF6FF",
        border_color="00267F"
    )

    # -------------------------------------------------------------
    # SECTION 4: DEMO NIGHT PRE-FLIGHT CHECKLIST
    # -------------------------------------------------------------
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Demo Night Pre-Flight Checklist (Quick 2-Minute Test)")
    r4.font.name = 'Arial'
    r4.font.size = Pt(15)
    r4.font.bold = True
    r4.font.color.rgb = NAVY
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)

    checklist = [
        ("Step 1: Access Platform", "Visit https://edudata-pmcup-app.surge.sh/ on laptop or iPad."),
        ("Step 2: Role Selection", "Select 'Statistician / Live Data Entry' -> Click 'Johnathan (johnathan)'."),
        ("Step 3: Enter Password", "Type 'password' and click 'Log in to Dashboard'."),
        ("Step 4: Sound & Alert Verification", "Click the green '🔔 Test Alert Signal' button in header. Confirm Fox 40 whistle sound plays."),
        ("Step 5: Inspect Scope Bar", "Verify that the Data Capturer Scope bar displays 'Master Lead Analyst (All Tiles)'."),
        ("Step 6: Live Fixture Launch", "Click any active fixture (e.g. Weymouth Wales vs. Paradise FC) to launch the live recording terminal.")
    ]

    for step_num, step_desc in checklist:
        p_chk = doc.add_paragraph()
        p_chk.paragraph_format.space_before = Pt(2)
        p_chk.paragraph_format.space_after = Pt(2)
        r_num = p_chk.add_run(f"✓  {step_num}: ")
        r_num.bold = True
        r_num.font.name = 'Calibri'
        r_num.font.size = Pt(10)
        r_num.font.color.rgb = EMERALD
        r_desc = p_chk.add_run(step_desc)
        r_desc.font.name = 'Calibri'
        r_desc.font.size = Pt(10)
        r_desc.font.color.rgb = SLATE

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_footer.add_run("— Prime Minister's Cup 2026 · Official Data Architecture & Operations —")
    r_foot.font.name = 'Calibri'
    r_foot.font.size = Pt(9)
    r_foot.font.italic = True
    r_foot.font.color.rgb = MUTED

    output_dir = r"c:\Users\noahb\OneDrive - The UWI - Cave Hill Campus\Desktop\EduData Project"
    artifact_dir = r"C:\Users\noahb\.gemini\antigravity\brain\e3ff380a-e042-4018-b625-5606c8fe8c48"
    
    file_name = "Prime_Ministers_Cup_Matchday_Operational_Guide.docx"
    doc_path = os.path.join(output_dir, file_name)
    art_path = os.path.join(artifact_dir, file_name)
    
    doc.save(doc_path)
    doc.save(art_path)
    print(f"Document successfully created at:\n1. {doc_path}\n2. {art_path}")

if __name__ == '__main__':
    create_document()
